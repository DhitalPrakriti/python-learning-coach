"""Render a project Markdown document to a styled .docx.

Usage:
    python tools/md_to_docx.py docs/SYSTEM_ARCHITECTURE.md out.docx \
        --title "Python Learning Coach" --subtitle "System Architecture"

Handles the Markdown subset these docs use: ATX headings, fenced code blocks,
pipe tables, bullet/numbered lists, blockquote callouts, horizontal rules as page
breaks, and inline **bold** / `code` spans. Also emits a title page, a generated
table of contents, and page numbers.

Deliberately not a general Markdown engine - it covers what docs/ actually
contains, and renders anything else as literal text rather than dropping it
silently.
"""

import argparse
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x55, 0x5F, 0x6D)
CODE_INK = RGBColor(0xA3, 0x1D, 0x4B)
CODE_BG = "F2F4F7"
QUOTE_BG = "FFF8ED"
QUOTE_EDGE = "B45309"

parser = argparse.ArgumentParser()
parser.add_argument("source")
parser.add_argument("output")
parser.add_argument("--title", default="Python Learning Coach")
parser.add_argument("--subtitle", default="Architecture Document")
args = parser.parse_args()


# --------------------------------------------------------------------------
# low-level docx helpers
# --------------------------------------------------------------------------
def shade(paragraph, fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(el)


def left_bar(paragraph, color):
    """Thick left border, for blockquote callouts."""
    borders = OxmlElement("w:pBdr")
    edge = OxmlElement("w:left")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), "18")
    edge.set(qn("w:space"), "8")
    edge.set(qn("w:color"), color)
    borders.append(edge)
    paragraph._p.get_or_add_pPr().append(borders)


def keep_together(paragraph):
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))


def add_field(paragraph, instr):
    """Insert a Word field code, used for PAGE / NUMPAGES."""
    r1 = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r1._r.append(begin)

    r2 = paragraph.add_run()
    txt = OxmlElement("w:instrText")
    txt.set(qn("xml:space"), "preserve")
    txt.text = instr
    r2._r.append(txt)

    r3 = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r3._r.append(end)


# --------------------------------------------------------------------------
# document setup
# --------------------------------------------------------------------------
doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.13

for name, size, before, after in (
    ("Heading 1", 22, 24, 12),
    ("Heading 2", 15, 16, 6),
    ("Heading 3", 12, 12, 4),
    ("Heading 4", 10.8, 10, 3),
):
    st = doc.styles[name]
    st.font.name = "Calibri Light"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = ACCENT
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

code_style = doc.styles.add_style("CodeBlock", 1)
code_style.font.name = "Consolas"
code_style.font.size = Pt(8.4)
code_style.paragraph_format.line_spacing = 1.0
code_style.paragraph_format.left_indent = Inches(0.14)

quote_style = doc.styles.add_style("Callout", 1)
quote_style.font.name = "Calibri"
quote_style.font.size = Pt(10)
quote_style.font.italic = True
quote_style.font.color.rgb = RGBColor(0x6B, 0x46, 0x0C)
quote_style.paragraph_format.left_indent = Inches(0.18)
quote_style.paragraph_format.space_before = Pt(6)
quote_style.paragraph_format.space_after = Pt(8)

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_field(footer, "PAGE")
footer.add_run(" of ")
add_field(footer, "NUMPAGES")
for r in footer.runs:
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED


def add_inline(paragraph, text, bold=False):
    """Emit runs, honouring **bold** and `code` spans."""
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.1)
            run.font.color.rgb = CODE_INK
        else:
            run = paragraph.add_run(part.replace("\\|", "|"))
        if bold:
            run.bold = True


def split_row(line):
    return [c.strip() for c in re.split(r"(?<!\\)\|", line.strip().strip("|"))]


lines = open(args.source, encoding="utf-8").read().split("\n")

# --------------------------------------------------------------------------
# title page
# --------------------------------------------------------------------------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(140)
r = p.add_run(args.title)
r.font.name = "Calibri Light"
r.font.size = Pt(34)
r.font.bold = True
r.font.color.rgb = ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(args.subtitle)
r.font.size = Pt(18)
r.font.color.rgb = MUTED

# The source opens with a `- **Key:** value` metadata block; put it on the cover.
meta = []
for line in lines[:16]:
    m = re.match(r"^- \*\*(.+?):\*\*\s*(.*)$", line.strip())
    if m:
        meta.append(f"{m.group(1)}: {re.sub(r'[`*]', '', m.group(2))}")

doc.add_paragraph()
for text in meta:
    q = doc.add_paragraph()
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    q.paragraph_format.space_after = Pt(3)
    rr = q.add_run(text)
    rr.font.size = Pt(10.5)
    rr.font.color.rgb = MUTED

doc.add_page_break()

# --------------------------------------------------------------------------
# contents (generated, so Word needs no field refresh)
# --------------------------------------------------------------------------
doc.add_paragraph(style="Heading 1").add_run("Contents")

seen_h1 = False
for line in lines:
    m = re.match(r"^(#{1,3})\s+(.*)$", line.strip())
    if not m:
        continue
    level, text = len(m.group(1)), re.sub(r"[`*]", "", m.group(2))
    if level == 1 and not seen_h1:
        seen_h1 = True  # document title, already on the cover
        continue
    e = doc.add_paragraph()
    e.paragraph_format.space_after = Pt(2)
    e.paragraph_format.left_indent = Inches({1: 0, 2: 0.22, 3: 0.46}[level])
    run = e.add_run(text)
    run.font.size = Pt(11 if level == 1 else 10)
    if level == 1:
        run.bold = True
        run.font.color.rgb = ACCENT
        e.paragraph_format.space_before = Pt(8)
    elif level == 3:
        run.font.color.rgb = MUTED
        run.font.size = Pt(9.5)

doc.add_page_break()

# --------------------------------------------------------------------------
# body
# --------------------------------------------------------------------------
i = 0
seen_h1 = False
skipped_meta = False

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # fenced code block
    if stripped.startswith("```"):
        i += 1
        body = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            body.append(lines[i])
            i += 1
        i += 1
        for n, code_line in enumerate(body):
            cp = doc.add_paragraph(style="CodeBlock")
            cp.paragraph_format.space_before = Pt(5 if n == 0 else 0)
            cp.paragraph_format.space_after = Pt(9 if n == len(body) - 1 else 0)
            cp.add_run(code_line if code_line.strip() else " ")
            shade(cp, CODE_BG)
        continue

    # pipe table
    if (
        stripped.startswith("|")
        and i + 1 < len(lines)
        and re.match(r"^\|[\s:\-|]+\|$", lines[i + 1].strip())
    ):
        header = split_row(stripped)
        i += 2
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append(split_row(lines[i]))
            i += 1

        table = doc.add_table(rows=1, cols=len(header))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell, text in zip(table.rows[0].cells, header):
            cell.paragraphs[0].text = ""
            add_inline(cell.paragraphs[0], text, bold=True)
        for row in rows:
            cells = table.add_row().cells
            for cell, text in zip(cells, row[: len(header)]):
                cell.paragraphs[0].text = ""
                add_inline(cell.paragraphs[0], text)
        for row in table.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    cp.paragraph_format.space_before = Pt(2)
                    cp.paragraph_format.space_after = Pt(2)
                    for r in cp.runs:
                        r.font.size = Pt(9)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        continue

    # blockquote callout, possibly spanning several lines
    if stripped.startswith(">"):
        chunk = []
        while i < len(lines) and lines[i].strip().startswith(">"):
            chunk.append(lines[i].strip().lstrip(">").strip())
            i += 1
        qp = doc.add_paragraph(style="Callout")
        add_inline(qp, " ".join(c for c in chunk if c))
        shade(qp, QUOTE_BG)
        left_bar(qp, QUOTE_EDGE)
        continue

    # horizontal rule -> page break
    if stripped in ("---", "***", "___"):
        doc.add_page_break()
        i += 1
        continue

    # headings
    m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
    if m:
        level, text = len(m.group(1)), m.group(2)
        if level == 1 and not seen_h1:
            seen_h1 = True  # already on the cover
            i += 1
            continue
        hp = doc.add_paragraph(style=f"Heading {level}")
        add_inline(hp, text)
        keep_together(hp)
        i += 1
        continue

    # the cover already carries the metadata block
    if not skipped_meta and re.match(r"^- \*\*(.+?):\*\*", stripped):
        while i < len(lines) and re.match(r"^- \*\*(.+?):\*\*", lines[i].strip()):
            i += 1
        skipped_meta = True
        continue

    # bullets
    m = re.match(r"^[-*]\s+(.*)$", stripped)
    if m:
        indented = len(line) - len(line.lstrip()) >= 2
        bp = doc.add_paragraph(style="List Bullet 2" if indented else "List Bullet")
        bp.paragraph_format.space_after = Pt(3)
        add_inline(bp, m.group(1))
        i += 1
        continue

    # numbered
    m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if m:
        np = doc.add_paragraph(style="List Number")
        np.paragraph_format.space_after = Pt(3)
        add_inline(np, m.group(2))
        i += 1
        continue

    # continuation line of a list item
    if line.startswith("   ") and stripped:
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Inches(0.5)
        cp.paragraph_format.space_after = Pt(3)
        add_inline(cp, stripped)
        i += 1
        continue

    if not stripped:
        i += 1
        continue

    bp = doc.add_paragraph()
    add_inline(bp, stripped)
    i += 1

doc.save(args.output)
print(f"wrote {args.output}")
